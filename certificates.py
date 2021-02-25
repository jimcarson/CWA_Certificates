import argparse
import datetime
import os
import sys
import pandas as pd
from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.shared import Cm, Inches, Pt, RGBColor
from docx2pdf import convert

DEBUG = False

#
# Defaults
#
ADVISOR_CALL = "WT8P"
ADVISOR_NAME = "Jim Carson"
ASSOCIATE_ADVISOR_CALL = None
ASSOCIATE_ADVISOR_NAME = None
CLASS_LEVEL = "Intermediate"
CLASS_LIST = None
CWA_SESSION_DATE = "Jan-Feb 2021"
OUTPUT_DIRECTORY = "output"
V_BOTTOM = WD_ALIGN_VERTICAL.BOTTOM

parser = argparse.ArgumentParser()
parser.add_argument("-c", "--classlist", type=str, help="File containing a list of students (call sign, name).  May also contain ClassLevel, AdvisorName, AdvisorCall, AssociateAdvisorName, AssociateAdvisorCall.", required=True)
parser.add_argument("-l", "--level", type=str, help="Level of class: Beginner, Basic, Intermediate or Advanced", required=False)
parser.add_argument("-d", "--date", type=str, help="Use a specific date for the report (Default: %s)" % CWA_SESSION_DATE, required=False)
parser.add_argument("-D", "--debug", help="Lots of print statements", action="store_true", required=False)
parser.add_argument("-o", "--outputdir", type=str, help="Output directory (Default: %s)" % OUTPUT_DIRECTORY, required=False)
parser.add_argument("-a", "--advisor", type=str, help="Advisor (Default: %s)" % ADVISOR_NAME, required=False)
parser.add_argument("-ac", "--advisor_call", type=str, help="Advisor Call Sign (Default; %s)" % ADVISOR_CALL, required=False)
parser.add_argument("-aa", "--associate_advisor", type=str, help="Associate Advisor (Default: %s)" % ASSOCIATE_ADVISOR_NAME, required=False)
parser.add_argument("-aac", "--associate_advisor_call", type=str, help="Associate Advisor Call Sign (Default: %s)" % ASSOCIATE_ADVISOR_CALL, required=False)
args = parser.parse_args()

if args.date:
    CWA_SESSION_DATE = args.date

if args.advisor:
    ADVISOR_NAME = args.advisor

if args.advisor_call:
    ADVISOR_CALL = args.advisor_call

if args.associate_advisor:
    ASSOCIATE_ADVISOR_NAME = args.associate_advisor

if args.associate_advisor_call:
    ASSOCIATE_ADVISOR_CALL = args.associate_advisor_call

if args.outputdir:
    OUTPUT_DIRECTORY = args.outputdir

# Create output directory if it doesn't already exist
if not os.path.exists(OUTPUT_DIRECTORY):
    os.makedirs(OUTPUT_DIRECTORY)

if args.level:
    CLASS_LEVEL = args.level

if args.debug:
    DEBUG = True

if args.classlist:
    try:
        CLASS_LIST = pd.read_csv(args.classlist, sep=",|\t", engine="python", index_col=0)
        if DEBUG:
           print(CLASS_LIST)
    except FileNotFoundError:
        print("File %s not found." % args.classlist)
        sys.exit(1)
    if "Name" not in CLASS_LIST.columns:
        print("Check the format of the class list.  The first two fields should be call sign and student name.")
        sys.exit(1)
else:
    print("Class list is required.")
    sys.exit(1)

if DEBUG:
    print("Output Directory: %s" % OUTPUT_DIRECTORY)
    print("CWA Class Date: %s" % CWA_SESSION_DATE)
    print("Class Level: %s" % CLASS_LEVEL)

for i in CLASS_LIST.index:
    name = CLASS_LIST.loc[i,"Name"]
    call = i
    if "ClassLevel" in CLASS_LIST.columns:
        CLASS_LEVEL = CLASS_LIST.loc[i,"ClassLevel"]
    if "AdvisorName" in CLASS_LIST.columns:
        ADVISOR_NAME = CLASS_LIST.loc[i,"AdvisorName"]
        ADVISOR_CALL = CLASS_LIST.loc[i,"AdvisorCall"]
    
    if "AssociateAdvisorName" in CLASS_LIST.columns:
        ASSOCIATE_ADVISOR_NAME = CLASS_LIST.loc[i,"AssociateAdvisorName"]
        ASSOCIATE_ADVISOR_CALL = CLASS_LIST.loc[i,"AssociateAdvisorCall"]

    if DEBUG:
        print("\t%s\t%s" %( call, name))
    D = Document("CWA_Template.docx")
    D._body.clear_content()
    """
    Add certificate title
    Add student name
    Add class level (beginner, basic, intermediate, advanced)
    Add table with:
        Left column = Advisor and optional Associate Advisor
        Center column = Term
        Right column = Kate and Joe, CW Academy Managers
    """
    p = D.add_paragraph(style="CWA_CERT_TITLE").add_run("CW Academy Certificate of Completion")
    p = D.add_paragraph(style="CWA_STUDENT_NAME").add_run("%s, %s" % (name, call))
    p = D.add_paragraph(style="CWA_NORMAL").add_run("has successfully completed")
    p = D.add_paragraph(style="CWA_CLASS_TYPE").add_run(CLASS_LEVEL)
    p = D.add_paragraph(style="CWA_NORMAL").add_run("an 8-week/16-session training program")
    p = D.add_paragraph(style="CWA_NORMAL").add_run("in Morse Code sending and receiving")

    # 
    # Now create the table at the bottom of the document.
    #
    t = D.add_table(rows=1, cols=3)
    t.allow_autofit = True

    (LC, CENTER_COL, RC) = t.add_row().cells

    # Left column - Advisor and optional Associate Advisor
    p = LC.add_paragraph(style="CWA_ADVISOR_SIG").add_run(ADVISOR_NAME)
    p = LC.add_paragraph(style="CWA_NORMAL").add_run("%s, %s\nCW Academy Advisor" % (ADVISOR_NAME, ADVISOR_CALL))
    if ASSOCIATE_ADVISOR_CALL:
        p = LC.add_paragraph(style="CWA_ASSOCIATE_ADVISOR_SIG").add_run(ASSOCIATE_ADVISOR_NAME)
        p = LC.add_paragraph(style="CWA_NORMAL").add_run("%s, %s\nCW Academy Assoc. Advisor" % (ASSOCIATE_ADVISOR_NAME, ASSOCIATE_ADVISOR_CALL))
    
    LC.vertical_alignment = V_BOTTOM
    
    # Center column - Class Date
    p = CENTER_COL.add_paragraph(style="CWA_NORMAL").add_run(CWA_SESSION_DATE)
    CENTER_COL.vertical_alignment = V_BOTTOM
    
    # Right column - CWA Managers
    p = RC.add_paragraph(style="CWA_KATE").add_run("Kate Hutton")
    p = RC.add_paragraph(style="CWA_NORMAL").add_run("Kate Hutton, K6HTN")
    p = RC.add_paragraph(style="CWA_ASSOCIATE_ADVISOR_SIG").add_run("Joe Fischer")
    p = RC.add_paragraph(style="CWA_NORMAL").add_run("Joe Fischer, AA8TA")
    p = RC.add_paragraph(style="CWA_NORMAL").add_run("CW Academy Managers")
    RC.vertical_alignment = V_BOTTOM

    # Save DOCX, then create a PDF from it.
    fn = os.path.join(OUTPUT_DIRECTORY,i+".docx")
    print("%s" % i)
    D.save(fn)
    convert(fn, fn.replace("docx","pdf"))
